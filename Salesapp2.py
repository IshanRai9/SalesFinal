import streamlit as st
import cohere
import PyPDF2
from docx import Document
from io import BytesIO
from docx.shared import Pt
import re
import pytesseract
from PIL import Image
from pdf2image import convert_from_bytes
from gmail_utils import gmail_authenticate, get_recent_emails, get_attachment
from dotenv import load_dotenv
import os
import json

# Write credentials.json to a temp file
with open("credentials.json", "w") as f:
    f.write(st.secrets["credentials"])

# Write token.json to a temp file
with open("token.json", "w") as f:
    f.write(st.secrets["token"])

cohere_api_key = st.secrets["cohere_api_key"]
co = cohere.Client(cohere_api_key)

data = None

# Set page configuration
st.set_page_config(page_title="Tender Summarizer", page_icon="📄")

# Custom CSS
st.markdown("""
    <style>

    /* General Layout */
    body {
        color: #495057;
        background-color: #f8f9fa;
        font-family: sans-serif;
    }

    .main .block-container {
        padding-bottom: 80px;
    }

    /* Header styling */
    .stApp > header {
        background-color: #0e1117;
        padding: 10px 20px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.04);
        position: sticky;
        top: 0;
        z-index: 1111;
    }

    /* Chat message styling */
    .chat-container {
        max-height: 600px;
        overflow-y: auto;
        padding: 20px;
        background-color: #ffffff;
        border-radius: 15px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
        margin-bottom: 20px;
    }

    .user-message {
        background: #e9ecef;
        font-weight: 500;
        border-radius: 8px;
        padding: 15px;
        margin-bottom: 10px;
        border: 1px solid #dee2e6;
        color: #495057;
        margin-left: auto;
        max-width: 85%;
    }

    .bot-message {
        background: #f8f9fa;
        border-radius: 8px;
        padding: 15px;
        margin-bottom: 10px;
        border: 1px solid #e9ecef;
        color: #495057;
        margin-right: auto;
        max-width: 85%;
    }

    /* Input styling */
    .stTextInput > div > div > input {
        border-radius: 8px;
        padding: 12px 16px;
        border: 2px solid #e9ecef;
        background-color: #fff;
        font-size: 14px;
        color: #495057;
        transition: all 0.2s ease;
    }

    .stTextInput > div > div > input:focus {
        border-color: #b22222;
        box-shadow: 0 0 0 3px rgba(178,34,34,0.1);
    }

    /* Button styling */
    .stButton > button {
        border-radius: 8px;
        padding: 10px 20px;
        background-color: #b22222 !important;
        color: white !important;
        border: none;
        cursor: pointer;
        box-shadow: 0 2px 3px rgba(0,0,0,0.1);
        transition: all 0.3s ease;
    }

    .stButton > button:hover {
        background-color: #8b0000 !important;
        transform: translateY(-1px);
    }

    /* Visit Us Button */
    .visit-button {
        display: inline-block;
        padding: 8px 20px;
        background-color: #b22222;
        color: white !important;
        text-decoration: none;
        border-radius: 6px;
        font-weight: 500;
        font-size: 14px;
        transition: all 0.3s ease;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        text-align: center;
        margin-right: 10px;
    }

    .visit-button:hover {
        background-color: #8b0000;
        transform: translateY(-1px);
        box-shadow: 0 4px 8px rgba(0,0,0,0.2);
        text-decoration: none;
    }

    /* File uploader styling */
    .stFileUploader > div > div > div {
        border-radius: 8px;
        border: 2px dashed #e9ecef;
        padding: 20px;
        background-color: #f9fafb;
    }

    /* Sidebar styling */
    .stSidebar > div:first-child {
        background-color: #f8f9fa;
        padding: 20px;
        border-radius: 8px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.05);
        z-index: 100;
    }

    /* Sidebar info styling */
    .sidebar-info {
        background-color: #ffffff;
        padding: 15px;
        border-radius: 8px;
        border: 1px solid #e9ecef;
        margin-bottom: 20px;
        font-size: 0.95em;
        color: #495057;
    }

    .sidebar-info h3 {
        font-size: 1.2em;
        color: #b22222;
        margin-bottom: 10px;
    }

    .sidebar-info ul {
        padding-left: 20px;
        margin: 0;
    }

    .sidebar-info li {
        margin-bottom: 8px;
    }
    </style>
""", unsafe_allow_html=True)

# Header with logo and visit button
with st.container():
    col1, col2, col3 = st.columns([1, 3, 1])
    with col1:
        st.image("https://s3ktech.ai/wp-content/uploads/2025/03/S3Ktech-Logo.png", width=140)
    with col2:
        pass  # No title
    with col3:
        st.markdown("""
            <div style="display: flex; justify-content: flex-end; align-items: center; height: 100%;">
                <a href="https://s3ktech.ai/" target="_blank" class="visit-button">Visit Us</a>
            </div>
        """, unsafe_allow_html=True)

# Sidebar with informational content
with st.sidebar:
    st.markdown("""
        <div class="sidebar-info">
            <h3>Welcome to Tender Summarizer</h3>
            <p>Streamline your tender analysis with AI-powered insights.</p>
            <ul>
                <li>Summarize tender emails</li>
                <li>Analyze attachments (PDF, DOCX, images)</li>
                <li>Download summaries as Word docs</li>
            </ul>
        </div>
    """, unsafe_allow_html=True)

def extract_text_from_scanned_pdf(pdf_file):
    images = convert_from_bytes(pdf_file.read())
    text = ""
    for i, image in enumerate(images):
        page_text = pytesseract.image_to_string(image, lang='eng')
        text += f"\n--- Page {i+1} ---\n{page_text}"
    return text

def extract_text_from_image(image_file):
    image = Image.open(image_file)
    return pytesseract.image_to_string(image, lang='eng')

def extract_text_from_pdf(pdf_file):
    reader = PyPDF2.PdfReader(pdf_file)
    return "".join(page.extract_text() or "" for page in reader.pages)

def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = "\n".join(para.text for para in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            text += "\n" + " ".join(cell.text for cell in row.cells)
    return text

def generate_table_word(summary_text):
    lines = summary_text.splitlines()
    heading = next((l.strip().lstrip('#').strip() for l in lines if l.strip().startswith('#')), 'Table')
    data = []
    key = None
    values = []
    i = 0
    while i < len(lines):
        line = lines[i].lstrip('#').strip()
        if re.match(r'^\*\*.*\*\*$', line):
            if key:
                data.append((key, values))
            key = line.strip('*').strip()
            values = []
        elif key and (line.startswith('-') or re.match(r'^\d+\.', line)):
            cleaned = re.sub(r'^[-\d.]+\s*', '', line)
            if cleaned not in values:
                values.append(cleaned)
        i += 1
    if key:
        data.append((key, values))
    doc = Document()
    title_para = doc.add_heading(level=1)
    run_title = title_para.add_run(heading)
    run_title.bold = True
    run_title.font.size = Pt(16)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Description'
    for key, vals in data:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = "\n".join(vals)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def stream_summary_from_cohere(text):
    global data
    x = ""
    prompt = (
        """You are an expert in analyzing and summarizing government and institutional tender documents.

            Summarize the following tender document by extracting and presenting all important and relevant information that may be present. Only include the sections that are explicitly mentioned or applicable in the document. **Do not include sections that are not present, not mentioned, or not relevant to the specific tender type.**

            **Make sure to use bullet points and headings to organize the information clearly and concisely.**
            
            1. AI Analyzed Lead Qualification and WIN Probability **MANDATORY**:
            -According to you Calculate the Lead Qualification (High Value Lead or Low Value Lead) and WIN Probability (High, Medium, Low) based on the information provided in the document.
            
            2. Extract details under the following categories, underline and increase the font by 2 of the categories compared to its description. **ONLY IF AVAILABLE ELSE DO NOT GIVE**:
            - Tender Name
            - Tender Reference Number and ID  
            - Name of the Issuing Organization or Authority  
            - Tender Fee (amount, mode of payment)  
            - EMD (Earnest Money Deposit) Details (amount, mode of payment)  
            - Estimated Tender Value or Project Cost  
            - Pre-bid Meeting Dates, Venue, and Registration/Link  
            - Tender Meeting Dates and Venues (if different from Pre-bid)  
            - Scope of Work  
            - Modules or Work Packages  
            - Workforce Requirements (specify onsite manpower and training manpower, if any)  
            - Human Resource Details  
            - Technical and Financial Eligibility Criteria  
            - Technical and Financial Marking/Scoring Criteria  
            - Performance Security Requirements  
            - Implementation Timeline and Phases (Turnaround Time or TAT)  
            - Contract Duration/Period  
            - Project Location(s)  
            - Existing IHMS or Software Application Details (if mentioned)  
            - Payment Terms and Schedule  
            - Submission Method (Online, Physical, or Hybrid)  
            - Selection Methodology (e.g., QCBS, L1)  
            - Cloud Service Provider (CSP) Details (if applicable)  
            - Hardware Details (especially for hospital/lab tenders—CT/MRI/X-ray/Pathology equipment)  
            - Technical Specifications  
            - Radiology/Pathology Scope (if applicable)  
            - Checklists (All the documents required, if provided)  
            - Declarations, Undertakings, and Affidavits  
            - OEM (Original Equipment Manufacturer) Document Requirements  
            - Penalty Clauses and Bidder Obligations  
            - Financial Bid Structure  
            - Viability Gap Funding (VGF)  
            - Special Purpose Vehicle (SPV) clauses  
            - Land Border Sharing Clause  
            - Mode of Payments for Tender Fee, EMD, and Other Charges  
            - Contact Details of the Tender Issuer (email, phone, address)
            Again, include **only the sections that are actually present in the document** and dont say not mentioned in the document, instead skip that section.

            3. Check for the following criteria based on the document: **COMPULSORY**
            -Mention if Criteria is yes or no or not able to find
            -Mention the line/statement to provide proof for "Yes" or "No" criteria.
            -If criteria is "not able to find", mention any similar statement that can be used as a proof if possible

                1.If the Make in India clause is yes, then we are not able to participate in tender.
                2.If the Average turnover  for last 3 Financial year is above is  30 Crore, then we cannot  participate. 
                3.If the GEM bid is in the BOQ category, then we can participate in the BID
                4.If the GEM Bid is in V2 Q2 category, then we cannot participate in the BID
                5.We have only ISO certififcate ISO9001 2015 and ISO20000 2028
                6.We have range of ADF Scanners from 20 PPM to 110 PM for A4 & A3 size documents
                7.Overhead Scanners up to A3 size

            Present the summary in a clean, organized format using clear headings or bullet points.
            
            4. At last give me these details seperate again: Tender Name, Tender Type (HIMS, Radiology Lab etc.), Tender registration start date and end date)
            Finally give a Final Summary using the information extracted from the document. The summary should be concise, clear, and easy to understand. It should provide a high-level overview of the tender document, highlighting the most important aspects without going into excessive detail. Make Sure to mention if the Bid is worth to chase or not along with the reason as to why it is worth or not worth to chase the bid.
            
            Tender Document:\n\n"""
        f"{text}"
    )
    response = co.chat_stream(
        model="command-a-03-2025",
        message=prompt
    )
    for chunk in response:
        if hasattr(chunk, "text") and chunk.text:
            yield chunk.text

def stream_email_summary_from_cohere(email_text, has_attachment):
    if not has_attachment:
        # Use tender-style prompt if no attachment is present
        prompt =(
            """You are an expert in analyzing and summarizing government and institutional tender emails.

            Summarize the following tender email by extracting and presenting all important and relevant information that may be present. Only include the sections that are explicitly mentioned or applicable in the email. **Do not include sections that are not present, not mentioned, or not relevant to the specific tender type.**

            **Make sure to use bullet points and headings to organize the information clearly and concisely.**
            
            1. AI Analyzed Lead Qualification and WIN Probability **MANDATORY**:
            -According to you Calculate the Lead Qualification (High Value Lead or Low Value Lead) and WIN Probability (High, Medium, Low) based on the information provided in the email.
            
            2. Extract details under the following categories, underline and increase the font by 2 of the categories compared to its description. **ONLY IF AVAILABLE ELSE DO NOT GIVE**:
            - Tender Name
            - Tender Reference Number and ID  
            - Name of the Issuing Organization or Authority  
            - Tender Fee (amount, mode of payment)  
            - EMD (Earnest Money Deposit) Details (amount, mode of payment)  
            - Estimated Tender Value or Project Cost  
            - Pre-bid Meeting Dates, Venue, and Registration/Link  
            - Tender Meeting Dates and Venues (if different from Pre-bid)  
            - Scope of Work  
            - Modules or Work Packages  
            - Workforce Requirements (specify onsite manpower and training manpower, if any)  
            - Human Resource Details  
            - Technical and Financial Eligibility Criteria  
            - Technical and Financial Marking/Scoring Criteria  
            - Performance Security Requirements  
            - Implementation Timeline and Phases (Turnaround Time or TAT)  
            - Contract Duration/Period  
            - Project Location(s)  
            - Existing IHMS or Software Application Details (if mentioned)  
            - Payment Terms and Schedule  
            - Submission Method (Online, Physical, or Hybrid)  
            - Selection Methodology (e.g., QCBS, L1)  
            - Cloud Service Provider (CSP) Details (if applicable)  
            - Hardware Details (especially for hospital/lab tenders—CT/MRI/X-ray/Pathology equipment)  
            - Technical Specifications  
            - Radiology/Pathology Scope (if applicable)  
            - Checklists (All the documents required, if provided)  
            - Declarations, Undertakings, and Affidavits  
            - OEM (Original Equipment Manufacturer) Document Requirements  
            - Penalty Clauses and Bidder Obligations  
            - Financial Bid Structure  
            - Viability Gap Funding (VGF)  
            - Special Purpose Vehicle (SPV) clauses  
            - Land Border Sharing Clause  
            - Mode of Payments for Tender Fee, EMD, and Other Charges  
            - Contact Details of the Tender Issuer (email, phone, address)
            Again, include **only the sections that are actually present in the email** and dont say not mentioned in the email, instead skip that section.

            3. Check for the following criteria based on the email: **COMPULSORY**
            -Mention if Criteria is yes or no or not able to find
            -Mention the line/statement to provide proof for "Yes" or "No" criteria.
            -If criteria is "not able to find", mention any similar statement that can be used as a proof if possible

                1.If the Make in India clause is yes, then we are not able to participate in tender.
                2.If the Average turnover  for last 3 Financial year is above is  30 Crore, then we cannot  participate. 
                3.If the GEM bid is in the BOQ category, then we can participate in the BID
                4.If the GEM Bid is in V2 Q2 category, then we cannot participate in the BID
                5.We have only ISO certififcate ISO9001 2015 and ISO20000 2028
                6.We have range of ADF Scanners from 20 PPM to 110 PM for A4 & A3 size documents
                7.Overhead Scanners up to A3 size

            Present the summary in a clean, organized format using clear headings or bullet points.
            
            4. At last give me these details seperate again: Tender Name, Tender Type (HIMS, Radiology Lab etc.), Tender registration start date and end date)
            Finally give a Final Summary using the information extracted from the email. The summary should be concise, clear, and easy to understand. It should provide a high-level overview of the tender email, highlighting the most important aspects without going into excessive detail. Make Sure to mention if the Bid is worth to chase or not along with the reason as to why it is worth or not worth to chase the bid.
            
            Tender Email:\n\n"""
            f"{email_text}"
        )
    else:
        # Use simple business email summary prompt
        prompt = (
            """Summarize the following business email in a clean and concise manner. 
            Highlight key intent, sender, and any mentioned attachments, deadlines, required actions or any *important details* in clean bullet points .\n\n"""
            f"Email:\n{email_text}"
        )

    response = co.chat_stream(
        model="command-a-03-2025",
        message=prompt
    )

    x = ""
    for chunk in response:
        if hasattr(chunk, "text") and chunk.text:
            yield chunk.text

# Streamlit App
st.title("📄 Tender Document Summarizer")

service = gmail_authenticate()
emails = get_recent_emails(service)

st.markdown("### ✉️ Recent Emails")
if emails:
    for email in emails:
        col1, col2, col3, col4 = st.columns([4, 3, 2, 2])
        col1.write(email['subject'])
        col2.write(email['from'])
        col3.write("✅ Yes" if email['has_attachment'] else "❌ No")
        if col4.button("Generate", key=email['id']):
            with st.spinner("Processing email and attachment..."):
                # 📧 Always show Email Summary
                email_text = email.get('snippet') or "No snippet available."
                email_summary_placeholder = st.empty()
                email_summary_output = ""
                for chunk in stream_email_summary_from_cohere(email_text, has_attachment=bool(email.get('has_attachment', False))):
                    email_summary_output += chunk
                    email_summary_placeholder.markdown("### 📖 Email Summary\n" + email_summary_output + "\n\n---\n")

                # 📎 Try to process the attachment (if any)
                filename, file_data = get_attachment(service, email['id'])
                if filename and file_data:
                    file_ext = filename.split('.')[-1].lower()
                    file_obj = BytesIO(file_data)

                    try:
                        if file_ext == 'pdf':
                            text = extract_text_from_pdf(file_obj)
                            if len(text.strip()) < 100:
                                file_obj.seek(0)
                                text = extract_text_from_scanned_pdf(file_obj)
                        elif file_ext == 'docx':
                            text = extract_text_from_docx(file_obj)
                        elif file_ext in ['png', 'jpg', 'jpeg', 'tiff']:
                            text = extract_text_from_image(file_obj)
                        else:
                            st.warning(f"Unsupported attachment type: {file_ext}")
                            text = ""

                        if text.strip():
                            summary_placeholder = st.empty()
                            summary_text = ""
                            for chunk in stream_summary_from_cohere(text):
                                summary_text += chunk
                                summary_placeholder.markdown("### 📄 Tender Summary\n" + summary_text)
                            st.session_state["summary"] = email_summary_output + summary_text

                            table_buffer = generate_table_word(email_summary_output + summary_text)
                            st.download_button(
                                "⬇️ Download Tender Summary",
                                data=table_buffer,
                                file_name=f"{filename}_summary.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        else:
                            st.warning("Could not extract valid text from the attachment.")
                    except Exception as e:
                        st.error(f"Error processing attachment: {str(e)}")
                else:
                    st.info("No attachment found. Email summary above.")

else:
    st.info("No recent emails found or authorized.")

st.markdown("---")
st.markdown("<p style='text-align:center; color: gray;'>Designed by Medimaze AI Team</p>", unsafe_allow_html=True)
